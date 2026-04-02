# /app/backend/routers/health.py
"""Health-check and monitoring endpoints."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel

import config
from backend.error_handler import APIError

try:
    import psutil
except ImportError:
    psutil = None


class FileModeRequest(BaseModel):
    mode: str = "local"
    base_url: Optional[str] = None
    source_root: Optional[str] = None


class PreviewExcelRequest(BaseModel):
    path: str
    page: int = 0
    page_size: int = 100


class CheckAccessRequest(BaseModel):
    path: str = ""

router = APIRouter(prefix="/api", tags=["health"])


@router.get("/health")
async def health_check():
    from backend.services.file_resolver import get_resolver
    resolver = get_resolver()
    return {
        "status": "ok",
        "engine": getattr(config, "ENGINE_NAME", "DevOps Analyzer"),
        "version": getattr(config, "ENGINE_VERSION", "unknown"),
        "file_mode": resolver.mode,
    }


@router.get("/file-mode")
async def get_file_mode():
    from backend.services.file_resolver import get_resolver
    return get_resolver().get_config()


@router.post("/file-mode")
async def set_file_mode(body: FileModeRequest):
    from backend.services.file_resolver import switch_mode
    kwargs = body.model_dump(exclude={"mode"}, exclude_none=True)
    resolver = switch_mode(body.mode, **kwargs)
    return {"ok": True, **resolver.get_config()}


@router.post("/preview-excel")
async def preview_excel_file(body: PreviewExcelRequest):
    """범용 문서 미리보기 — 로컬 경로에서 시트별 데이터 반환"""
    file_path = body.path.strip()
    page = body.page
    page_size = body.page_size
    if not file_path:
        raise APIError(status_code=400, message="path required", code="MISSING_PATH")
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        raise HTTPException(status_code=404, detail=f"파일 없음: {file_path}")
    row_start = page * page_size
    row_end = row_start + page_size

    ext = p.suffix.lower()

    try:
        if ext in ('.xlsx', '.xls', '.xlsm'):
            import openpyxl
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
            sheets = []

            for name in wb.sheetnames:
                ws = wb[name]
                # Find actual header row (first row with multiple non-empty values)
                header_row = 0
                all_rows_raw = []
                for ri, row in enumerate(ws.iter_rows(values_only=True)):
                    str_row = [str(c or '') for c in row]
                    all_rows_raw.append(str_row)
                    if ri >= 200:
                        break

                # For test spec sheets, find the row with TC ID / Test Case ID
                is_test_spec = any(kw in name.lower() for kw in ['test spec', 'test case', 'traceability', 'unit test'])
                if is_test_spec:
                    # Find the most detailed header row (with TC ID, Title, etc.)
                    best_row = 0
                    best_score = 0
                    for ri, row in enumerate(all_rows_raw[:15]):
                        non_empty = sum(1 for c in row if c.strip())
                        has_id = any('id' in c.lower() for c in row)
                        has_detail = any(kw in ' '.join(row).lower() for kw in ['title', 'description', 'method', 'environment', 'result', 'function'])
                        score = non_empty + (3 if has_id else 0) + (3 if has_detail else 0)
                        if score > best_score:
                            best_score = score
                            best_row = ri
                    header_row = best_row

                headers = all_rows_raw[header_row] if header_row < len(all_rows_raw) else []
                while headers and not headers[-1].strip():
                    headers.pop()
                all_data = all_rows_raw[header_row + 1:]
                data_rows = [r[:len(headers)] for r in all_data[row_start:row_end]]

                sheets.append({
                    "name": name,
                    "headers": headers,
                    "rows": data_rows,
                    "total_rows": ws.max_row - header_row,
                    "total_cols": len(headers),
                })

            wb.close()
            # Keep only sheets with meaningful data — skip Cover, History, Introduction
            skip = {'cover', 'history'}
            useful = [s for s in sheets
                      if s["name"].lower() not in skip
                      and len([r for r in s.get("rows", []) if any(c.strip() for c in r)]) >= 2]
            return {"ok": True, "filename": p.name, "sheets": useful if useful else sheets, "sheet_names": [s["name"] for s in (useful if useful else sheets)]}

        elif ext == '.docx':
            import docx as _docx
            import re as _re
            doc = _docx.Document(str(p))

            # Detect structured tables (UDS Function Info / SDS Component Info)
            func_tables = []
            comp_tables = []
            attr_tables = []
            other_tables = []
            for i, table in enumerate(doc.tables):
                first_text = table.rows[0].cells[0].text.strip() if table.rows else ""
                if "Software Component" in first_text and "Information" in first_text:
                    # SDS component table — parse as key-value
                    comp_data = {}
                    for row in table.rows[1:]:
                        cells = [c.text.strip() for c in row.cells]
                        if len(cells) >= 3:
                            label = cells[0] or cells[1]
                            value = cells[2]
                            if label and "Software Component" not in label:
                                comp_data[label] = value
                    if comp_data.get("ID") or comp_data.get("Name") or comp_data.get("SC ID") or comp_data.get("SC Name"):
                        comp_tables.append(comp_data)
                elif first_text == "Attribute" and len(table.rows) <= 10:
                    # SDS attribute table
                    attr_data = {}
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        if len(cells) >= 2 and cells[0]:
                            attr_data[cells[0]] = cells[1] if len(cells) > 1 else ""
                    if attr_data:
                        attr_tables.append(attr_data)
                elif "Function Information" in first_text:
                    func_data = {}
                    ns_a = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
                    for row in table.rows[1:]:
                        cells_raw = row.cells
                        cells = [c.text.strip() for c in cells_raw]
                        if len(cells) >= 3:
                            label = cells[0] or cells[1]
                            value = cells[2]
                            if label and label != "[ Function Information ]":
                                func_data[label] = value
                            if label == "Logic Diagram":
                                for cell in cells_raw[2:]:
                                    blips = cell._element.findall('.//a:blip', ns_a)
                                    for b in blips:
                                        embed = b.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                        if embed:
                                            func_data["_image_id"] = embed
                                            break
                                    if func_data.get("_image_id"):
                                        break
                    if func_data.get("ID") or func_data.get("Name"):
                        func_tables.append(func_data)
                else:
                    rows_data = []
                    headers = []
                    for ri, row in enumerate(table.rows):
                        # Deduplicate merged cells
                        raw = [c.text.strip() for c in row.cells]
                        deduped = []
                        prev = None
                        for c in raw:
                            if c != prev:
                                deduped.append(c)
                            prev = c
                        if ri == 0:
                            headers = deduped
                        else:
                            rows_data.append(deduped)
                        if ri >= 100:
                            break
                    if headers or rows_data:
                        other_tables.append({
                            "name": f"Table {len(other_tables)+1}",
                            "headers": headers,
                            "rows": rows_data[:100],
                            "total_rows": len(table.rows),
                            "total_cols": len(headers),
                        })

            sheets = []

            # SDS Component list
            if comp_tables:
                comp_keys = ["SC ID", "SC Name", "ID", "Name", "Description", "ASIL", "Related ID",
                             "Allocated Requirements", "Allocated Function", "Sub-Components", "Interface"]
                comp_headers = [k for k in comp_keys if any(c.get(k) for c in comp_tables)]
                # Add meaningful extra keys (skip numeric, bracket-prefixed)
                extra = sorted(k for k in {k for c in comp_tables for k in c.keys()} - set(comp_headers)
                               if k and not k.isdigit() and not k.startswith("[") and not k.startswith("N/A") and len(k) > 2)
                comp_headers.extend(extra[:10])  # limit extras
                comp_rows = [[c.get(k, "") for k in comp_headers] for c in comp_tables[:200]]
                sheets.append({
                    "name": f"Components ({len(comp_tables)})",
                    "headers": comp_headers,
                    "rows": comp_rows[:100],
                    "total_rows": len(comp_tables),
                    "total_cols": len(comp_headers),
                })

            # SDS Attribute list
            if attr_tables:
                attr_keys = sorted({k for a in attr_tables for k in a.keys()})
                attr_rows = [[a.get(k, "") for k in attr_keys] for a in attr_tables[:200]]
                sheets.append({
                    "name": f"Attributes ({len(attr_tables)})",
                    "headers": attr_keys,
                    "rows": attr_rows[:100],
                    "total_rows": len(attr_tables),
                    "total_cols": len(attr_keys),
                })

            # UDS Function list as structured table
            if func_tables:
                func_keys = ["ID", "Name", "Prototype", "Description", "ASIL", "Related ID",
                              "Input Parameters", "Output Parameters", "Called Function", "Calling Function"]
                func_headers = [k for k in func_keys if any(f.get(k) for f in func_tables)]
                # Add Logic Diagram column if images exist
                has_images = any(f.get("_image_id") for f in func_tables)
                if has_images:
                    func_headers.append("Logic Diagram")
                func_rows = []
                for f in func_tables[row_start:row_end]:
                    row_data = [f.get(k, "") for k in func_headers[:-1] if k != "Logic Diagram"]
                    if has_images:
                        img_id = f.get("_image_id", "")
                        row_data.append(f"__IMG__{img_id}" if img_id else "")
                    func_rows.append(row_data)
                sheets.append({
                    "name": f"Functions ({len(func_tables)})",
                    "headers": func_headers,
                    "rows": func_rows[:100],
                    "total_rows": len(func_tables),
                    "total_cols": len(func_headers),
                })

            # Other tables
            sheets.extend(other_tables[:10])

            # Paragraphs
            paras = [pg.text for pg in doc.paragraphs if pg.text.strip()][:200]
            sheets.append({
                "name": "Content",
                "headers": ["Text"],
                "rows": [[t] for t in paras[:100]],
                "total_rows": len(paras),
                "total_cols": 1,
            })

            # Keep only named sheets (Functions, Components, Attributes, Content) — remove generic Table N
            useful = [s for s in sheets if not s["name"].startswith("Table ")]
            return {"ok": True, "filename": p.name, "sheets": useful if useful else sheets[:3], "sheet_names": [s["name"] for s in (useful if useful else sheets[:3])]}

        elif ext == '.txt':
            text = p.read_text(encoding='utf-8', errors='replace')
            lines = text.splitlines()[:200]
            return {"ok": True, "filename": p.name, "sheets": [{
                "name": "Content",
                "headers": ["Line"],
                "rows": [[l] for l in lines[:100]],
                "total_rows": len(lines),
                "total_cols": 1,
            }], "sheet_names": ["Content"]}

        else:
            raise HTTPException(status_code=400, detail=f"지원하지 않는 형식: {ext}")

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@router.get("/preview-image")
async def preview_image(path: str, image_id: str):
    """docx 문서에서 이미지 추출 반환"""
    from fastapi.responses import Response
    p = Path(path).expanduser().resolve()
    if not p.exists():
        raise HTTPException(status_code=404)
    try:
        import docx as _docx
        doc = _docx.Document(str(p))
        rel = doc.part.rels.get(image_id)
        if not rel or 'image' not in rel.reltype:
            raise HTTPException(status_code=404, detail="image not found")
        blob = rel.target_part.blob
        ct = rel.target_part.content_type or 'image/png'
        return Response(content=blob, media_type=ct)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/file-mode/check-access")
async def check_cloudium_access(body: CheckAccessRequest = CheckAccessRequest()):
    """경로 접근 가능 여부 확인."""
    from backend.services.file_resolver import get_resolver
    resolver = get_resolver()
    test_path = body.path
    if test_path:
        try:
            resolver._check_allowed(test_path) if hasattr(resolver, '_check_allowed') else None
            accessible = Path(test_path).exists()
            return {"ok": True, "accessible": accessible, "path": test_path, "mode": resolver.mode}
        except PermissionError as e:
            return {"ok": False, "accessible": False, "error": str(e), "mode": resolver.mode}
    return {"ok": True, "mode": resolver.mode, **resolver.get_config()}


@router.get("/metrics")
async def metrics():
    """Detailed system metrics for monitoring."""
    if psutil is None:
        return {
            "cpu_percent": None,
            "memory": {"total_mb": None, "used_percent": None},
            "disk": {"free_gb": None},
            "process": {"pid": os.getpid(), "threads": None},
            "note": "psutil not installed — install for full metrics",
        }
    return {
        "cpu_percent": psutil.cpu_percent(interval=0.1),
        "memory": {
            "total_mb": psutil.virtual_memory().total // (1024 * 1024),
            "used_percent": psutil.virtual_memory().percent,
        },
        "disk": {
            "free_gb": psutil.disk_usage("/").free // (1024**3),
        },
        "process": {
            "pid": os.getpid(),
            "threads": len(psutil.Process(os.getpid()).threads()),
        },
    }


@router.post("/cache/clear")
async def clear_cache():
    """Clear all in-memory caches."""
    from backend import state
    state.jenkins_progress.clear()
    state.uds_view_cache.clear()
    state.source_sections_cache.clear()
    state.session_list_cache.clear()
    return {"ok": True, "message": "All caches cleared"}
